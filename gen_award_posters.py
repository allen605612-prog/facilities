#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
得獎海報產生器 v2.1 — A1 300 DPI 版
用法：
  uv run --with pillow --with python-docx --with numpy --with img2pdf gen_award_posters.py <docx路徑> [輸出資料夾]
"""
import sys, os, re, math, random, argparse
import numpy as np
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
import docx as _docx

# ── A1 @ 300 DPI ─────────────────────────────────────────────
_BASE_W, _BASE_H = 800, 1200       # 原設計基準（僅用來算比例）
DPI = 300
A1_W_MM, A1_H_MM = 594, 841
W = int(A1_W_MM * DPI / 25.4)     # 7016 px
H = int(A1_H_MM * DPI / 25.4)     # 9933 px
SX = W / _BASE_W                   # ≈ 8.77  — 水平縮放
SY = H / _BASE_H                   # ≈ 8.28  — 垂直縮放

def sx(v):  return int(v * SX)
def sy(v):  return int(v * SY)
def sf(v):  return max(1, int(v * SX))   # 字型大小（以水平比例為準）

# ── 字型 ─────────────────────────────────────────────────────
_FONT_PATHS = [
    r"C:\Windows\Fonts\msjh.ttc",
    r"C:\Windows\Fonts\msjhbd.ttc",
    r"C:\Windows\Fonts\kaiu.ttf",
    r"C:\Windows\Fonts\mingliu.ttc",
]

def _load_font(size: int) -> ImageFont.FreeTypeFont:
    for fp in _FONT_PATHS:
        try:
            return ImageFont.truetype(fp, size)
        except Exception:
            continue
    return ImageFont.load_default()

# ── 色盤 ─────────────────────────────────────────────────────
NAVY       = (10,  18,  40)
NAVY_MID   = (18,  30,  65)
GOLD       = (212, 175,  55)
GOLD_LIGHT = (255, 223, 100)
GOLD_DARK  = (160, 120,  20)
WHITE      = (255, 255, 255)
OFF_WHITE  = (240, 230, 200)


# ══════════════════════════════════════════════════════════════
# 1. 解析 docx
# ══════════════════════════════════════════════════════════════

def _collect_lines(path: Path) -> list:
    doc = _docx.Document(str(path))
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in lines:
                    lines.append(t)
    return lines


def _parse_event_info(title: str) -> tuple:
    m = re.search(r'第(\d+)屆', title)
    year = m.group(1) if m else '?'
    clean = title.lstrip('「」『』')
    rm = re.match(r'(.{2,6}?)第\d+屆', clean)
    region = rm.group(1) if rm else ''
    if '國民中小學' in title:
        etype = '國民中小學科展'
    elif '分區' in title and '高中' in title:
        etype = '分區高中科展'
    elif '高中' in title:
        etype = '高中科展'
    elif '國中' in title:
        etype = '國中科展'
    else:
        etype = '科學展覽會'
    return year, region, etype


def _parse_students(s: str) -> list:
    groups = []
    current_cls = ''
    current_names = []
    for part in s.split('、'):
        part = part.strip().replace('　', ' ')
        if not part:
            continue
        if ' ' in part:
            if current_cls:
                groups.append((current_cls, current_names))
            seg = part.split(' ', 1)
            current_cls = seg[0]
            current_names = [seg[1]] if len(seg) > 1 else []
        else:
            current_names.append(part)
    if current_cls:
        groups.append((current_cls, current_names))
    return groups


def parse_docx(path: Path) -> tuple:
    lines = _collect_lines(path)
    title = ''
    for line in lines:
        if re.search(r'第\d+屆', line):
            title = line.strip('「」『』')
            break
    year, region, etype = _parse_event_info(title)

    awards = []
    cur = {}
    rank_pat = re.compile(r'^(.+?)\s+(第[一二三四五六七八九十百]+名)\s*$')
    stu_pat  = re.compile(r'^參加學生[：:]\s*(.+)$')
    tch_pat  = re.compile(r'^指導老師[：:]\s*(.+?)(?:\s*老師)?\s*$')

    for line in lines:
        m = rank_pat.match(line)
        if m and not line.startswith(('參', '指', '感')):
            if cur.get('subject'):
                awards.append(cur)
            cur = {'subject': m.group(1).strip(), 'rank': m.group(2).strip()}
            continue
        m = stu_pat.match(line)
        if m:
            cur['names'] = _parse_students(m.group(1).strip())
            continue
        m = tch_pat.match(line)
        if m:
            teacher = re.sub(r'\s*老師\s*$', '', m.group(1).strip())
            cur['teacher'] = teacher
            continue

    if cur.get('subject'):
        awards.append(cur)

    return title, year, region, etype, awards


# ══════════════════════════════════════════════════════════════
# 2. 繪圖工具
# ══════════════════════════════════════════════════════════════

def _text_w(draw, text, font):
    bb = draw.textbbox((0, 0), text, font=font)
    return bb[2] - bb[0]

def _centered(draw, text, y, font, fill):
    x = (W - _text_w(draw, text, font)) // 2
    draw.text((x, y), text, font=font, fill=fill)

def _shadow(draw, text, y, font, fill, shd=(0, 0, 0)):
    x = (W - _text_w(draw, text, font)) // 2
    off = max(2, sx(2))
    draw.text((x + off, y + off), text, font=font, fill=shd)
    draw.text((x, y), text, font=font, fill=fill)

def _auto_font(draw, text, max_w, max_sz, min_sz=sf(26)):
    for sz in range(max_sz, min_sz - 1, -2):
        f = _load_font(sz)
        if _text_w(draw, text, f) <= max_w:
            return f
    return _load_font(min_sz)


def _draw_bg(img, draw):
    """numpy 加速版漸層背景"""
    arr = np.zeros((H, W, 3), dtype=np.float32)
    t = np.linspace(0, 1, H)
    sin_t = np.sin(t * np.pi)
    for ch in range(3):
        row_vals = NAVY[ch] + (NAVY_MID[ch] - NAVY[ch]) * sin_t
        arr[:, :, ch] = row_vals[:, np.newaxis]

    # 垂直條紋光影
    stripe_hw = max(1, sx(32))
    for i in range(16):
        cx = int(W * i / 15)
        bright = 28.0 if i % 2 == 0 else -14.0
        x_start = max(0, cx - stripe_hw)
        x_end   = min(W, cx + stripe_hw)
        xs = np.arange(x_start, x_end)
        fade = bright * (1.0 - np.abs(xs - cx) / stripe_hw)
        arr[:, x_start:x_end, :] = np.clip(
            arr[:, x_start:x_end, :] + fade[np.newaxis, :, np.newaxis], 0, 255
        )

    img.paste(Image.fromarray(arr.astype(np.uint8), 'RGB'))


def _draw_confetti(draw, seed):
    rng = random.Random(seed)
    colors = [GOLD, GOLD_LIGHT, (255,200,50), (255,245,160), (200,150,30), (220,180,60)]
    for _ in range(130):
        x = rng.randint(0, W)
        y = rng.randint(int(H * .08), int(H * .88))
        c = rng.choice(colors)
        kind = rng.randint(0, 3)
        if kind == 0:
            ang = rng.uniform(-70, 70)
            l  = rng.randint(sx(8), sx(20))
            w2 = rng.randint(sx(2), sx(6))
            pts = [(-l/2, -w2/2), (l/2, -w2/2), (l/2, w2/2), (-l/2, w2/2)]
            ca, sa = math.cos(math.radians(ang)), math.sin(math.radians(ang))
            draw.polygon([(x + p[0]*ca - p[1]*sa, y + p[0]*sa + p[1]*ca) for p in pts], fill=c)
        elif kind == 1:
            r2 = rng.randint(sx(2), sx(6))
            draw.ellipse([x-r2, y-r2, x+r2, y+r2], fill=c)
        elif kind == 2:
            x2 = x + rng.randint(-sx(18), sx(18))
            y2 = y + rng.randint(sy(4), sy(22))
            draw.line([x, y, x2, y2], fill=c, width=max(1, sx(2)))
        else:
            s_sz = rng.randint(sx(4), sx(9))
            draw.polygon([(x, y-s_sz), (x+s_sz, y), (x, y+s_sz), (x-s_sz, y)], fill=c)


def _star_pts(cx, cy, ro, ri, n=5, off=-90):
    return [(cx + (ro if i%2==0 else ri) * math.cos(math.radians(off + i*180/n)),
             cy + (ro if i%2==0 else ri) * math.sin(math.radians(off + i*180/n)))
            for i in range(n*2)]

def _draw_star(draw, cx, cy, ro, ri):
    draw.polygon(_star_pts(cx, cy, ro, ri), fill=GOLD_DARK)
    draw.polygon(_star_pts(cx - ro*.07, cy - ro*.07, ro*.84, ri*.84), fill=GOLD)
    draw.polygon(_star_pts(cx - ro*.13, cy - ro*.13, ro*.54, ri*.54), fill=GOLD_LIGHT)


def _draw_bar(draw, y):
    draw.rectangle([0, y,        W, y + sy(56)], fill=GOLD_DARK)
    draw.rectangle([0, y + sy(4), W, y + sy(52)], fill=GOLD)
    draw.line([(0, y + sy(4)),  (W, y + sy(4))],  fill=GOLD_LIGHT, width=max(1, sx(3)))
    draw.line([(0, y + sy(52)), (W, y + sy(52))], fill=GOLD_DARK,  width=max(1, sx(2)))


def _vignette(img):
    steps = min(sx(120), min(W, H) // 2)
    vig = Image.new("L", (W, H), 255)
    vd  = ImageDraw.Draw(vig)
    for i in range(steps):
        vd.rectangle([i, i, W-i, H-i], outline=int(255 * i / steps))
    return Image.composite(img, Image.new("RGB", (W, H), 0), vig)


# ══════════════════════════════════════════════════════════════
# 3. 名字排版
# ══════════════════════════════════════════════════════════════

def _format_name_lines(names_data):
    parts = [cls + ' ' + '、'.join(ns) for cls, ns in names_data]
    full  = '、'.join(parts)
    if len(full) <= 18:
        return [full + ' 同學']
    if len(parts) == 1:
        ns_all = names_data[0][1]
        cls    = names_data[0][0]
        mid    = len(ns_all) // 2 + (len(ns_all) % 2)
        return [cls + ' ' + '、'.join(ns_all[:mid]),
                '、'.join(ns_all[mid:]) + ' 同學']
    mid = len(parts) // 2 + (1 if len(parts) % 2 else 0)
    return ['、'.join(parts[:mid]), '、'.join(parts[mid:]) + ' 同學']


# ══════════════════════════════════════════════════════════════
# 4. 海報主體（A1 300 DPI）
# ══════════════════════════════════════════════════════════════

def make_poster(award, year, region, event_type, out_path, seed=42):
    img  = Image.new("RGB", (W, H), NAVY)
    draw = ImageDraw.Draw(img)
    _draw_bg(img, draw)

    draw = ImageDraw.Draw(img)   # 貼完背景後重建 draw
    _draw_confetti(draw, seed)

    _draw_star(draw, sx(290), sy(305), sx(55), sx(22))
    _draw_star(draw, sx(510), sy(310), sx(60), sx(24))
    _draw_star(draw, sx(400), sy(245), sx(80), sx(32))

    f_crazy  = _load_font(sf(54))
    f_names  = _load_font(sf(36))
    f_achiev = _load_font(sf(56))
    f_thanks = _load_font(sf(36))
    MAX_W    = W - sx(80)

    _shadow(draw, "狂 賀", sy(390), f_crazy, GOLD_LIGHT)

    name_lines = _format_name_lines(award.get('names', []))
    y_n = sy(468)
    for line in name_lines:
        f_nl = _auto_font(draw, line, MAX_W, sf(36), sf(26))
        _centered(draw, line, y_n, f_nl, OFF_WHITE)
        y_n += sy(54)

    y_div = y_n + sy(4)
    draw.line([(sx(80), y_div), (W - sx(80), y_div)], fill=GOLD_DARK, width=max(1, sx(1)))

    subj   = award.get('subject', '')
    rank   = award.get('rank', '')
    line_a1 = f"榮獲第 {year} 屆{region}"
    line_a2 = event_type
    line_a3 = f"「{subj}」{rank}！"

    if _text_w(draw, line_a3, f_achiev) <= MAX_W:
        y_a, lh = y_div + sy(148), sy(70)
        _shadow(draw, line_a1,     y_a,        f_achiev, GOLD_LIGHT, (80, 60, 0))
        _shadow(draw, line_a2,     y_a + lh,   f_achiev, GOLD_LIGHT, (80, 60, 0))
        _shadow(draw, line_a3,     y_a + lh*2, f_achiev, GOLD_LIGHT, (80, 60, 0))
        y_last = y_a + lh * 2
    else:
        y_a, lh = y_div + sy(100), sy(63)
        subj_line = f"「{subj}」"
        f_subj = _auto_font(draw, subj_line, MAX_W, sf(56), sf(28))
        _shadow(draw, line_a1,      y_a,        f_achiev, GOLD_LIGHT, (80, 60, 0))
        _shadow(draw, line_a2,      y_a + lh,   f_achiev, GOLD_LIGHT, (80, 60, 0))
        _shadow(draw, subj_line,    y_a + lh*2, f_subj,   GOLD_LIGHT, (80, 60, 0))
        _shadow(draw, f"{rank}！",  y_a + lh*3, f_achiev, GOLD_LIGHT, (80, 60, 0))
        y_last = y_a + lh * 3

    bar_y = min(y_last + sy(95), H - sy(198))
    _draw_bar(draw, bar_y)
    _centered(draw, f"感謝 {award.get('teacher', '')} 老師指導",
              bar_y + sy(88), f_thanks, WHITE)

    img = _vignette(img)
    img.save(str(out_path), "PNG", dpi=(DPI, DPI))
    print(f"  PNG: {out_path}  ({W}×{H} @ {DPI} DPI)")

    # 個別 PDF
    pdf_path = out_path.with_suffix('.pdf')
    try:
        import img2pdf
        layout = img2pdf.get_layout_fun(
            (img2pdf.mm_to_pt(A1_W_MM), img2pdf.mm_to_pt(A1_H_MM))
        )
        with open(pdf_path, 'wb') as f:
            f.write(img2pdf.convert(str(out_path), layout_fun=layout))
        print(f"  PDF: {pdf_path}")
    except ImportError:
        pass   # img2pdf 未安裝時跳過


# ══════════════════════════════════════════════════════════════
# 5. 主程式
# ══════════════════════════════════════════════════════════════

def main():
    ap = argparse.ArgumentParser(description='從 docx 批次產生 A1 得獎海報')
    ap.add_argument('docx',   help='Word 文件路徑（.docx）')
    ap.add_argument('outdir', nargs='?', default=None, help='輸出資料夾（預設：文件同資料夾）')
    args = ap.parse_args()

    docx_path = Path(args.docx)
    if not docx_path.exists():
        print(f"錯誤：找不到檔案 {docx_path}", file=sys.stderr)
        sys.exit(1)

    outdir = Path(args.outdir) if args.outdir else docx_path.parent
    outdir.mkdir(parents=True, exist_ok=True)

    print(f"輸出尺寸：{W}×{H} px（A1 @ {DPI} DPI）")
    title, year, region, etype, awards = parse_docx(docx_path)
    print(f"活動：{title}")
    print(f"屆次：第 {year} 屆 | 地區：{region} | 類型：{etype}")
    print(f"獎項：{len(awards)} 項\n")

    out_pngs = []
    for i, award in enumerate(awards, 1):
        fname    = f"poster_{year}_{i:02d}_{award['subject']}{award['rank']}.png"
        out_path = outdir / fname
        print(f"[{i:02d}/{len(awards):02d}] {award['subject']} {award['rank']}")
        make_poster(award, year, region, etype, out_path, seed=40+i)
        out_pngs.append(str(out_path))

    # 合併 PDF
    if out_pngs:
        try:
            import img2pdf
            merged = outdir / f"海報集_第{year}屆_A1_300dpi.pdf"
            layout = img2pdf.get_layout_fun(
                (img2pdf.mm_to_pt(A1_W_MM), img2pdf.mm_to_pt(A1_H_MM))
            )
            with open(merged, 'wb') as f:
                f.write(img2pdf.convert(out_pngs, layout_fun=layout))
            print(f"\n合併 PDF: {merged}")
        except ImportError:
            pass

    print(f"\n完成！共 {len(awards)} 張，儲存於：{outdir}")


if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    main()
