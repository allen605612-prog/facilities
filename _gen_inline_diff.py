"""
比對 設備組.docx vs 高中資訊內控制度_完整版v5.docx
輸出 inline diff：刪除→灰色刪除線，新增/修改→紅色粗體
"""
import sys, copy, difflib
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX

GRAY = RGBColor(0x80, 0x80, 0x80)
RED  = RGBColor(0xC0, 0x00, 0x00)
FONT = '標楷體'

# ── 載入兩份文件 ──────────────────────────────────────────────
doc1 = Document('D:/D/114設備組/設備組.docx')
doc2 = Document('D:/D/114設備組/高中資訊內控制度_完整版v5.docx')

# ── 取得 body 內所有段落（含表格）元素 ───────────────────────
# 回傳 list of (type, text, xml_element)
def body_items(doc):
    items = []
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            from docx.text.paragraph import Paragraph
            p = Paragraph(child, doc)
            items.append(('p', p.text, child, p))
        elif tag == 'tbl':
            # 取表格所有文字做 key
            texts = []
            for row in child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                for t in row.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    if t.text:
                        texts.append(t.text)
            items.append(('tbl', ' '.join(texts), child, None))
    return items

items1 = body_items(doc1)
items2 = body_items(doc2)
texts1 = [it[1] for it in items1]
texts2 = [it[1] for it in items2]

# ── 輸出文件 ──────────────────────────────────────────────────
out = Document()
# 套用基本字型
style = out.styles['Normal']
style.font.name = FONT
style.font.size = Pt(11)

def set_font_xml(run_elem, name):
    rpr = run_elem.find(qn('w:rPr'))
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        run_elem.insert(0, rpr)
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), name)
    rf.set(qn('w:hAnsi'), name)
    rf.set(qn('w:eastAsia'), name)
    rpr.insert(0, rf)

def apply_del_style(para_obj):
    """對整段每個 run 套用灰色刪除線"""
    for run in para_obj.runs:
        run.font.color.rgb = GRAY
        run.font.strike = True
        if not run.font.name:
            run.font.name = FONT

def apply_red_style(para_obj):
    """對整段每個 run 套用紅色粗體；若無 run 則新建一個"""
    if not para_obj.runs:
        run = para_obj.add_run(para_obj.text)
        para_obj.clear()
        para_obj.add_run(para_obj.text or '')
    for run in para_obj.runs:
        run.font.color.rgb = RED
        run.font.bold = True
        run.font.name = FONT

def add_label(out_doc, label, color):
    """在輸出文件加一個說明段落"""
    p = out_doc.add_paragraph()
    run = p.add_run(label)
    run.font.color.rgb = color
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = FONT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def copy_elem_del(out_doc, item):
    """複製元素到輸出文件，套用刪除線樣式"""
    kind, text, xml_elem, para_obj = item
    if kind == 'p':
        # 複製整個 <w:p> XML
        new_p = copy.deepcopy(xml_elem)
        out_doc.element.body.append(new_p)
        # 套用刪除線到所有 run
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        for run in new_p.findall(f'{{{ns}}}r'):
            rpr = run.find(f'{{{ns}}}rPr')
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run.insert(0, rpr)
            # color
            col = OxmlElement('w:color')
            col.set(qn('w:val'), '808080')
            rpr.append(col)
            # strike
            strike = OxmlElement('w:strike')
            rpr.append(strike)
    elif kind == 'tbl':
        # 複製表格並在每個 run 加刪除線
        new_tbl = copy.deepcopy(xml_elem)
        out_doc.element.body.append(new_tbl)
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        for run in new_tbl.findall(f'.//{{{ns}}}r'):
            rpr = run.find(f'{{{ns}}}rPr')
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run.insert(0, rpr)
            col = OxmlElement('w:color')
            col.set(qn('w:val'), '808080')
            rpr.append(col)
            strike = OxmlElement('w:strike')
            rpr.append(strike)

def copy_elem_red(out_doc, item):
    """複製元素到輸出文件，套用紅色粗體"""
    kind, text, xml_elem, para_obj = item
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    if kind == 'p':
        new_p = copy.deepcopy(xml_elem)
        out_doc.element.body.append(new_p)
        for run in new_p.findall(f'{{{ns}}}r'):
            rpr = run.find(f'{{{ns}}}rPr')
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run.insert(0, rpr)
            col = OxmlElement('w:color')
            col.set(qn('w:val'), 'C00000')
            rpr.append(col)
            b = OxmlElement('w:b')
            rpr.append(b)
    elif kind == 'tbl':
        new_tbl = copy.deepcopy(xml_elem)
        out_doc.element.body.append(new_tbl)
        for run in new_tbl.findall(f'.//{{{ns}}}r'):
            rpr = run.find(f'{{{ns}}}rPr')
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run.insert(0, rpr)
            col = OxmlElement('w:color')
            col.set(qn('w:val'), 'C00000')
            rpr.append(col)
            b = OxmlElement('w:b')
            rpr.append(b)

def copy_elem_same(out_doc, item):
    """直接複製元素（不改樣式）"""
    kind, text, xml_elem, para_obj = item
    out_doc.element.body.append(copy.deepcopy(xml_elem))

# ── 執行 diff ─────────────────────────────────────────────────
out_body = out.element.body

matcher = difflib.SequenceMatcher(None, texts1, texts2, autojunk=False)
opcodes = matcher.get_opcodes()

print(f'opcodes: {len(opcodes)}')
for tag, i1, i2, j1, j2 in opcodes:
    print(f'  {tag:8s} doc1[{i1}:{i2}] doc2[{j1}:{j2}]')

for tag, i1, i2, j1, j2 in opcodes:
    if tag == 'equal':
        for i in range(i1, i2):
            copy_elem_same(out, items1[i])

    elif tag == 'delete':
        for i in range(i1, i2):
            copy_elem_del(out, items1[i])

    elif tag == 'insert':
        for j in range(j1, j2):
            copy_elem_red(out, items2[j])

    elif tag == 'replace':
        # 先顯示舊內容（刪除線），再顯示新內容（紅字）
        for i in range(i1, i2):
            copy_elem_del(out, items1[i])
        for j in range(j1, j2):
            copy_elem_red(out, items2[j])

# ── 移除最後一個多餘的空 body 子元素（sectPr 位置問題）───────
# 把原本 out 的 sectPr 保留，避免頁面設定錯亂
outpath = 'D:/D/114設備組/高中資訊內控制度_比對版.docx'
out.save(outpath)
print('saved:', outpath)
