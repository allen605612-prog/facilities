"""
產生兩份混合題型試卷（docx）
第一節：教育理念選擇題 + 雙語教學英文知能 + 申論題
第二節：教育理念申論題 + 國語文選擇題 + 數學填充題
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "標楷體"

def set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

def _run(para, text, size=12, bold=False, italic=False):
    r = para.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold   = bold
    r.italic = italic
    # 設定東亞字型
    rPr = r._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT)
    rPr.insert(0, rFonts)
    return r

def heading(doc, text, size=14, center=True, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    _run(p, text, size=size, bold=bold)
    return p

def body(doc, text, size=11, indent=0.0, bold=False, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent   = Cm(indent)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(2)
    _run(p, text, size=size, bold=bold)
    return p

def question(doc, num, q_text, options, size=11, en=False):
    """出一道選擇題"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(1)
    _run(p, f"{num}. ", size=size, bold=True)
    _run(p, q_text, size=size)
    for opt in options:
        p_opt = doc.add_paragraph()
        p_opt.paragraph_format.left_indent  = Cm(1.5)
        p_opt.paragraph_format.space_before = Pt(0)
        p_opt.paragraph_format.space_after  = Pt(0)
        _run(p_opt, opt, size=size)

def essay_q(doc, num_str, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    _run(p, f"{num_str}、", size=size, bold=True)
    _run(p, text, size=size)

def note_item(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    _run(p, text, size=size)

# ─────────────────────────────────────────────
# ★ 第一節試卷
# ─────────────────────────────────────────────
def build_paper1(path):
    doc = Document()
    set_margins(doc)

    heading(doc, "國立臺中教育大學", 15)
    heading(doc, "國民小學教師及加註雙語教學次專長學士後教育學分班（B 班）招生考試試題", 13)
    heading(doc, "考試科目：教師意向（教育理念）暨雙語教學知能測驗　　　考試時間：1 小時", 12)
    doc.add_paragraph()

    body(doc, "注意事項：", 11, bold=True)
    note_item(doc, "一、本試題共三大部分，總計 100 分。請依題號順序作答至答案卷，於本試題卷作答者不予計分。")
    note_item(doc, "二、第一、二部分為單一選擇題，複選作答者該題不予計分；答案卡請用 2B 鉛筆清楚劃記，答錯不倒扣。")
    note_item(doc, "三、第三部分為申論題，請以黑、藍色原子筆或鋼筆填寫於答案卷上。")
    note_item(doc, "四、禁止使用電子計算機或任何儀器。")
    doc.add_paragraph()

    # ── 第一部分：教育理念中文選擇題 ──────────────────────
    body(doc, "第一部分：教育理念選擇題（共 10 題，每題 4 分，共 40 分）", 12, bold=True)
    doc.add_paragraph()

    p1s1 = [
        (1,  "「學習共同體」強調學生彼此對話合作、教室對外開放共同觀摩，此一教育理念起源自哪個國家？",
             ["（A）美國", "（B）英國", "（C）日本", "（D）芬蘭"]),
        (2,  "Peters 所提出的教育三大規準中，「合價值性」是指教育內容須符合哪項條件？",
             ["（A）具有認知意義且能增進知識",
              "（B）值得學習且對學習者具有正面意義",
              "（C）學習者出於自願、非被迫接受",
              "（D）符合社會法律規範與制度"]),
        (3,  "十二年國民基本教育課程綱要以「終身學習者」為主軸，下列何者為其三大核心素養面向之一？",
             ["（A）知識記憶", "（B）技能演練", "（C）自主行動", "（D）服從指令"]),
        (4,  "Vygotsky「近側發展區」（ZPD）最主要的教學意涵為何？",
             ["（A）教師應只教學生已經會的概念，確保成功學習經驗",
              "（B）教師需提供鷹架支持，協助學生跨越現有能力邊界",
              "（C）學習應完全由學生自主探索，教師不介入",
              "（D）標準化測驗是評估 ZPD 最有效的工具"]),
        (5,  "班杜拉（Bandura）社會學習論的「觀察學習」，依序包含哪四個歷程？",
             ["（A）注意→保留→再生→動機",
              "（B）刺激→反應→增強→消弱",
              "（C）感知→理解→應用→評鑑",
              "（D）探索→發現→驗證→歸納"]),
        (6,  "差異化教學（Differentiated Instruction）的核心，在於教師依學生「準備度」調整教學。下列何種做法最能體現此理念？",
             ["（A）全班使用同一份測驗，一律不做任何調整",
              "（B）針對不同程度的學生設計難度相異的學習任務",
              "（C）只關注學習落後的學生，忽視資優學生的需求",
              "（D）採用同儕競爭排名激勵全體學生"]),
        (7,  "依 Bloom 修訂版認知目標分類，「能比較並評斷不同教育政策的優缺點」屬於哪一層次？",
             ["（A）記憶（Remember）",
              "（B）理解（Understand）",
              "（C）分析（Analyze）",
              "（D）評鑑（Evaluate）"]),
        (8,  "Gardner 多元智能理論中，擅長辨別音調、旋律與節奏的智能稱為何者？",
             ["（A）語文智能", "（B）音樂智能", "（C）空間智能", "（D）自然觀察智能"]),
        (9,  "Seligman PERMA 幸福模型的五大要素中，「R」代表哪個要素？",
             ["（A）韌性（Resilience）",
              "（B）獎勵（Reward）",
              "（C）人際關係（Relationships）",
              "（D）宗教信仰（Religion）"]),
        (10, "教師在學校展現「組織公民行為」（Organizational Citizenship Behavior），下列何種行為最能代表此概念？",
             ["（A）依規定完成授課時數後即刻離校",
              "（B）主動協助新進同事備課、參與校務且不要求報酬",
              "（C）只有在校長明確要求時才參與課程研討",
              "（D）依薪資高低決定投入教學的程度"]),
    ]
    for num, q, opts in p1s1:
        question(doc, num, q, opts)
        doc.add_paragraph()

    # ── 第二部分：英文雙語教學知能 ─────────────────────────
    doc.add_page_break()
    body(doc, "第二部分：雙語教學英文知能選擇題（共 10 題，每題 4 分，共 40 分）", 12, bold=True)
    doc.add_paragraph()

    p1s2 = [
        (11, "Which of the following best describes 'additive bilingualism'?",
             ["(A) A process in which learning a second language gradually replaces the learner's first language.",
              "(B) A situation in which a second language enriches the learner's repertoire without displacing the first.",
              "(C) A condition in which two languages merge into a single hybrid communication system.",
              "(D) The simultaneous acquisition of two languages before the age of three."]),
        (12, "In the CLIL (Content and Language Integrated Learning) framework, which of the '4Cs' refers to using language to develop higher-order thinking skills?",
             ["(A) Content", "(B) Communication", "(C) Cognition", "(D) Culture"]),
        (13, "According to Krashen's Input Hypothesis, language acquisition is best promoted when input is at level 'i+1'. This means the input should be:",
             ["(A) Far beyond the learner's current comprehension level.",
              "(B) Exactly at the learner's current level with no new elements.",
              "(C) Slightly above the learner's current level of competence.",
              "(D) Simplified to well below the learner's level."]),
        (14, "Cummins distinguished between BICS and CALP. Which of the following is an example of CALP (Cognitive Academic Language Proficiency)?",
             ["(A) Chatting with classmates during recess.",
              "(B) Ordering food in the school cafeteria.",
              "(C) Writing an analytical essay on a historical event.",
              "(D) Greeting a teacher in the hallway."]),
        (15, "Which of the following correctly defines 'translanguaging' in bilingual education?",
             ["(A) Translating academic texts word-for-word from one language to another.",
              "(B) A pedagogical practice that treats the bilingual learner's full linguistic repertoire as a resource.",
              "(C) The requirement for teachers to use only the target language in all instruction.",
              "(D) A method for testing which language a student is stronger in."]),
        (16, "In Kachru's 'Three Circles' model of World Englishes, which group of countries belongs to the 'outer circle'?",
             ["(A) USA, UK, and Australia",
              "(B) China, Japan, and South Korea",
              "(C) India, Singapore, and Nigeria",
              "(D) Brazil, Mexico, and Argentina"]),
        (17, "Which bilingual education model explicitly aims to develop full bilingualism and biliteracy in both a minority and a majority language?",
             ["(A) Submersion program",
              "(B) Transitional bilingual program",
              "(C) Dual language immersion program",
              "(D) English-only immersion program"]),
        (18, "In the CLIL 'language triptych', which type of language refers to unexpected vocabulary and structures that emerge naturally during content learning?",
             ["(A) Language of learning — technical vocabulary specific to the subject.",
              "(B) Language for learning — language needed to manage classroom tasks.",
              "(C) Language through learning — language that emerges unexpectedly during the learning process.",
              "(D) Language for assessment — language used to evaluate student understanding."]),
    ]
    for num, q, opts in p1s2:
        question(doc, num, q, opts, en=True)
        doc.add_paragraph()

    # Reading passage Q19–20
    p_r = doc.add_paragraph()
    p_r.paragraph_format.left_indent = Cm(0.5)
    p_r.paragraph_format.space_after = Pt(4)
    _run(p_r,
         "Read the following passage and answer questions 19–20.\n\n"
         "    Research into bilingual individuals has consistently shown that managing two languages "
         "may confer certain cognitive advantages. Bilingual speakers must constantly select the appropriate "
         "language and inhibit the other, which engages executive control functions. This habitual cognitive "
         "exercise has been associated with enhanced attention, better task-switching abilities, and, in some "
         "studies, a delayed onset of cognitive decline in older adults. However, the extent of these benefits "
         "remains debated, as factors such as the degree of bilingualism, frequency of language use, and "
         "socioeconomic background may all influence outcomes. Critics caution that overgeneralizing bilingual "
         "advantages risks overlooking the complex realities of language learning.",
         size=11, italic=True)
    doc.add_paragraph()

    question(doc, 19,
             "According to the passage, which cognitive ability is most directly associated with bilingualism?",
             ["(A) Improved long-term memory for factual information.",
              "(B) Enhanced executive control, including attention and task-switching.",
              "(C) Superior spatial reasoning and mathematical ability.",
              "(D) Faster first-language acquisition in early childhood."])
    doc.add_paragraph()

    question(doc, 20,
             "What does the passage suggest regarding the debate over the cognitive advantages of bilingualism?",
             ["(A) The cognitive benefits of bilingualism are universal and apply equally to all individuals.",
              "(B) All research studies consistently replicate the same bilingual cognitive advantages.",
              "(C) The degree of benefit may vary due to factors such as frequency of language use and socioeconomic background.",
              "(D) Critics fully dismiss the existence of any cognitive benefits associated with bilingualism."])
    doc.add_paragraph()

    # ── 第三部分：申論題 ────────────────────────────────────
    doc.add_page_break()
    body(doc, "第三部分：申論題（共 1 題，20 分）", 12, bold=True)
    body(doc, "※作答內容請勿書寫服務單位、姓名或可供辨識身分之個人資料；違反者，將逕依情節酌以扣分。",
         10.5, indent=0.5)
    doc.add_paragraph()

    essay_q(doc, "一",
            "教育部積極推動「2030 雙語國家政策」，鼓勵各縣市國民小學逐步推展雙語教育，"
            "以英語作為部分學科之教學媒介語。然而，部分學者認為此政策可能輕忽母語文化傳承，"
            "並對學習資源較為不足的偏鄉學校造成更大的教育不平等。"
            "試就正反兩方觀點，評析雙語教育政策對國民小學學生學習與文化認同之潛在影響，"
            "並提出您認為在國小推動雙語教育時應把握的關鍵原則。（20 分）")

    doc.add_paragraph()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_end, "— 試題結束，請仔細檢查後繳交！—", size=11, bold=True)

    doc.save(path)
    print(f"✓ 第一節試卷已儲存：{path}")


# ─────────────────────────────────────────────
# ★ 第二節試卷
# ─────────────────────────────────────────────
def build_paper2(path):
    doc = Document()
    set_margins(doc)

    heading(doc, "國立臺中教育大學", 15)
    heading(doc, "國民小學教師及加註雙語教學次專長學士後教育學分班（B 班）招生考試試題", 13)
    heading(doc, "考試科目：教師意向測驗（國語文暨數學基本能力）　　　考試時間：1 小時", 12)
    doc.add_paragraph()

    body(doc, "注意事項：", 11, bold=True)
    note_item(doc, "一、本試題共三大部分，總計 100 分。請依題號順序作答至答案卷，於本試題卷作答者不予計分。")
    note_item(doc, "二、第二部分為單一選擇題，複選作答者該題不予計分；答案卡請用 2B 鉛筆清楚劃記，答錯不倒扣。")
    note_item(doc, "三、第一、三部分請以黑、藍色原子筆或鋼筆作答至答案卷，不須抄題，但請標明題號。")
    note_item(doc, "四、禁止使用電子計算機或任何儀器。")
    doc.add_paragraph()

    # ── 第一部分：教育理念申論題 ─────────────────────────
    body(doc, "第一部分：教育理念申論題（共 2 題，每題 20 分，共 40 分）", 12, bold=True)
    body(doc, "※作答內容請勿書寫服務單位、姓名或可供辨識身分之個人資料；違反者，將逕依情節酌以扣分。",
         10.5, indent=0.5)
    doc.add_paragraph()

    essay_q(doc, "一",
            "非洲有句諺語：「養育一個孩子，需要整個村落的力量。」（It takes a village to raise a child.）"
            "在當前多元社會中，教師不再是學生學習的唯一資源提供者。試從「親師合作」與「社區資源連結」"
            "兩個面向，說明教師如何善用外部夥伴關係，促進學生全人發展與學習成效。（20 分）")
    doc.add_paragraph()

    essay_q(doc, "二",
            "「素養導向教學」是十二年國教課程改革的核心精神之一。"
            "試說明「素養」的意涵為何，並就「真實情境」、「跨領域整合」與「學習遷移」三個面向，"
            "闡述教師在設計素養導向課程與教學時，應掌握的原則及具體做法。（20 分）")
    doc.add_paragraph()

    # ── 第二部分：國語文基本能力選擇題 ──────────────────
    doc.add_page_break()
    body(doc, "第二部分：國語文基本能力選擇題（共 10 題，每題 3 分，共 30 分）", 12, bold=True)
    doc.add_paragraph()

    p2s2 = [
        (1,  "下列何者字形完全正確？",
             ["（A）出類拔「萃」、草「菅」人命",
              "（B）義無「返」顧、汗牛充「棟」",
              "（C）一「劍」鍾情、氣勢「洶洶」",
              "（D）沉「緬」往事、鼎力「相」助"]),
        (2,  "「天時不如地利，地利不如人和」運用了何種修辭法？",
             ["（A）對偶、排比",
              "（B）頂真、層遞",
              "（C）類疊、映襯",
              "（D）譬喻、誇飾"]),
        (3,  "下列四個句子，何者無錯別字？",
             ["（A）他面對困境仍「不棄不撓」，展現出令人欽佩的毅力。",
              "（B）他「力排眾議」，堅持己見，終於得到師長肯定。",
              "（C）演講者引「徵」博引，滔滔不絕，令聽眾折服。",
              "（D）她的文章立「義」不夠深刻，有待加強。"]),
        (4,  "「鳥宿池邊樹，僧敲月下門」中，詩人反覆斟酌「敲」字的選用，後世因此演變出哪個成語？",
             ["（A）字斟句酌",
              "（B）推陳出新",
              "（C）推敲（反覆斟酌文字之意）",
              "（D）一字之師"]),
        (5,  "下列句子，何者為「存現句」？",
             ["（A）他昨天跑去圖書館借書。",
              "（B）山腳下散落著幾間農舍。",
              "（C）多讀書可以增廣見聞。",
              "（D）她是我最敬重的老師。"]),
        (6,  "「白璧微瑕」中，「璧」字的造字方法屬於哪一種？",
             ["（A）象形", "（B）指事", "（C）會意", "（D）形聲"]),
        (7,  "「東邊日出西邊雨，道是無晴卻有晴」同時運用了哪兩種修辭？",
             ["（A）誇飾、對偶",
              "（B）雙關、對偶",
              "（C）排比、映襯",
              "（D）頂真、譬喻"]),
        (8,  "「衍聲複詞」（聯綿詞）的特點是兩個音節共同表達一個意義，不可分割解釋。下列何者不是聯綿詞？",
             ["（A）蝴蝶", "（B）彷彿", "（C）蜘蛛", "（D）草木"]),
        (9,  "下列何者語用最為適當？",
             ["（A）「陳伯伯「享年」八十歲，至今仍耳聰目明、身強體健。」",
              "（B）「敝人拙作奉上，敬請老師「惠賜」指正。」",
              "（C）「請您「拜讀」此篇拙文，並賜予寶貴意見。」",
              "（D）「您的「拙作」已拜讀完畢，深感佩服，謹此致謝。」"]),
        (10, "李白〈靜夜思〉：「床前明月光，疑是地上霜，舉頭望明月，低頭思故鄉。」所表達的主要情感為何？",
             ["（A）傷時感事", "（B）思鄉懷遠", "（C）傷懷弔古", "（D）閨怨相思"]),
    ]
    for num, q, opts in p2s2:
        question(doc, num, q, opts)
        doc.add_paragraph()

    # ── 第三部分：數學基本能力填充題 ────────────────────
    doc.add_page_break()
    body(doc, "第三部分：數學基本能力填充題（共 3 題，每題 10 分，共 30 分）", 12, bold=True)
    body(doc, "（請以黑、藍色原子筆或鋼筆，依序將答案填入答案卷的空格中）", 10.5, indent=0.5)
    doc.add_paragraph()

    math_qs = [
        "1.　數線上 A、B 兩點分別代表 −6 與 10。若 C 點在 A、B 之間且 AC：CB＝1：3，則 C 點所代表的數為（　　　　　）。",
        "2.　已知等差數列的第 3 項為 11，第 7 項為 27，求此數列前 12 項之和為（　　　　　）。",
        "3.　若方程式 2x² − 8x + k ＝ 0 恰有兩個相等實數根，求 k 的值為（　　　　　）。",
    ]
    for mq in math_qs:
        p_mq = doc.add_paragraph()
        p_mq.paragraph_format.space_before = Pt(6)
        p_mq.paragraph_format.space_after  = Pt(6)
        _run(p_mq, mq, size=11)
        doc.add_paragraph()

    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_end, "— 試題結束，請仔細檢查後繳交！—", size=11, bold=True)

    doc.save(path)
    print(f"✓ 第二節試卷已儲存：{path}")


if __name__ == "__main__":
    build_paper1(r"D:\D\onedrive\文件\新版試卷_第一節_教師意向暨雙語教學知能.docx")
    build_paper2(r"D:\D\onedrive\文件\新版試卷_第二節_教師意向測驗暨國語文數學.docx")
    print("\n✅ 兩份試卷均已產生完畢！")
