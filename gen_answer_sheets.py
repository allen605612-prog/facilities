"""
產生兩份答案卷（docx）
第一節：選擇題答案表 × 2 + 申論題書寫區
第二節：申論題書寫區 × 2 + 選擇題答案表 + 數學填充答案表
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "標楷體"


# ── 基礎工具 ────────────────────────────────────────────────

def set_margins(doc, top=1.8, bottom=1.8, left=2.5, right=2.5):
    for sec in doc.sections:
        sec.top_margin    = Cm(top)
        sec.bottom_margin = Cm(bottom)
        sec.left_margin   = Cm(left)
        sec.right_margin  = Cm(right)


def _run(para, text, size=12, bold=False):
    r = para.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    rPr = r._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT)
    rPr.insert(0, rFonts)
    return r


def blk(doc, text='', size=12, bold=False, center=False,
        indent=0.0, sb=2, sa=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent   = Cm(indent)
    p.paragraph_format.space_before  = Pt(sb)
    p.paragraph_format.space_after   = Pt(sa)
    if text:
        _run(p, text, size=size, bold=bold)
    return p


def set_cell_border(cell, sides):
    """sides: dict like {'bottom': 'single', 'top': 'none', ...}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove old tcBorders if any
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in sides.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), val)
        if val != 'none':
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_row_height(row, cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(Cm(cm).emu / 914.4)))  # emu→twips
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def cell_center(cell, text, size=11, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    _run(p, text, size=size, bold=bold)


# ── 元件 ────────────────────────────────────────────────────

def add_header(doc, subtitle, subject, time_str):
    """試卷頂部標題 + 個人資料欄"""
    blk(doc, '國立臺中教育大學', 15, bold=True, center=True, sb=0, sa=2)
    blk(doc, subtitle, 13, bold=True, center=True, sb=0, sa=2)
    blk(doc, '招生考試答案卷', 13, bold=True, center=True, sb=0, sa=4)

    # 個人資料表
    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = 'Table Grid'

    labels_r1 = ['准考證號碼', '姓　　名', '座位號碼']
    for i, lbl in enumerate(labels_r1):
        cell_center(tbl.cell(0, i), lbl, size=11, bold=True)
    for i in range(3):
        tbl.cell(1, i).text = ''
        tbl.cell(1, i).paragraphs[0].paragraph_format.space_before = Pt(0)
        tbl.cell(1, i).paragraphs[0].paragraph_format.space_after  = Pt(0)
    for row in tbl.rows:
        set_row_height(row, 1.0)

    blk(doc, f'考試科目：{subject}　　考試時間：{time_str}',
        11, center=True, sb=4, sa=6)


def add_mc_answer_grid(doc, label, start, end):
    """橫向選擇題答案格（題號一列，答案一列）"""
    blk(doc, label, 11, bold=True, sb=8, sa=4)

    nums = list(range(start, end + 1))
    ncols = len(nums) + 1          # +1 for label column

    tbl = doc.add_table(rows=2, cols=ncols)
    tbl.style = 'Table Grid'

    # 標題列
    cell_center(tbl.cell(0, 0), '題　號', 10, bold=True)
    for i, n in enumerate(nums):
        cell_center(tbl.cell(0, i + 1), str(n), 10)

    # 答案列
    cell_center(tbl.cell(1, 0), '答　案', 10, bold=True)
    for i in range(len(nums)):
        tbl.cell(1, i + 1).text = ''
        tbl.cell(1, i + 1).paragraphs[0].paragraph_format.space_before = Pt(0)
        tbl.cell(1, i + 1).paragraphs[0].paragraph_format.space_after  = Pt(0)

    for row in tbl.rows:
        set_row_height(row, 0.9)


def add_essay_space(doc, label, n_lines=28):
    """申論題書寫區（有底線的空白行）"""
    blk(doc, label, 11, bold=True, sb=10, sa=6)

    for _ in range(n_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

        # 段落底線
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        btm = OxmlElement('w:bottom')
        btm.set(qn('w:val'), 'single')
        btm.set(qn('w:sz'), '4')
        btm.set(qn('w:space'), '1')
        btm.set(qn('w:color'), '888888')
        pBdr.append(btm)
        pPr.append(pBdr)

        r = p.add_run('　')   # 全形空白給足行高
        r.font.size = Pt(20)


def add_math_answer_boxes(doc, label, q_labels):
    """數學填充題答案格"""
    blk(doc, label, 11, bold=True, sb=8, sa=4)

    tbl = doc.add_table(rows=2, cols=len(q_labels) + 1)
    tbl.style = 'Table Grid'

    cell_center(tbl.cell(0, 0), '題號', 10, bold=True)
    for i, ql in enumerate(q_labels):
        cell_center(tbl.cell(0, i + 1), ql, 10)

    cell_center(tbl.cell(1, 0), '答案', 10, bold=True)
    for i in range(len(q_labels)):
        tbl.cell(1, i + 1).text = ''
        tbl.cell(1, i + 1).paragraphs[0].paragraph_format.space_before = Pt(0)
        tbl.cell(1, i + 1).paragraphs[0].paragraph_format.space_after  = Pt(0)

    set_row_height(tbl.rows[0], 0.9)
    set_row_height(tbl.rows[1], 1.4)


# ─────────────────────────────────────────────────────────────
# ★ 第一節答案卷
# ─────────────────────────────────────────────────────────────
def build_answer1(path):
    doc = Document()
    set_margins(doc)

    add_header(
        doc,
        '國民小學教師及加註雙語教學次專長學士後教育學分班（B 班）',
        '教師意向（教育理念）暨雙語教學知能測驗',
        '1 小時'
    )

    blk(doc, '【注意】本答案卷共三部分，請依題號作答；申論題請以黑、藍色原子筆或鋼筆填寫。',
        10, indent=0.3, sb=0, sa=4)

    # 第一部分：選擇題 Q1–10
    add_mc_answer_grid(doc,
        '■ 第一部分：教育理念選擇題（Q1–10，每題 4 分，共 40 分）', 1, 10)

    # 第二部分：選擇題 Q11–20
    add_mc_answer_grid(doc,
        '■ 第二部分：雙語教學英文知能選擇題（Q11–20，每題 4 分，共 40 分）', 11, 20)

    # 第三部分：申論題
    doc.add_page_break()
    blk(doc, '■ 第三部分：申論題（共 1 題，20 分）', 11, bold=True, sb=0, sa=4)
    blk(doc, '※ 作答內容請勿書寫服務單位、姓名或可供辨識身分之個人資料。', 10, indent=0.3, sb=0, sa=6)

    add_essay_space(doc, '第一題（20 分）', n_lines=35)

    # 頁尾
    blk(doc, '', sb=10, sa=0)
    p_end = blk(doc, '— 答案卷結束 —', 11, bold=True, center=True)

    doc.save(path)
    print(f'✓ 第一節答案卷：{path}')


# ─────────────────────────────────────────────────────────────
# ★ 第二節答案卷
# ─────────────────────────────────────────────────────────────
def build_answer2(path):
    doc = Document()
    set_margins(doc)

    add_header(
        doc,
        '國民小學教師及加註雙語教學次專長學士後教育學分班（B 班）',
        '教師意向測驗（國語文暨數學基本能力）',
        '1 小時'
    )

    blk(doc, '【注意】本答案卷共三部分，請依題號作答；申論題、填充題請以黑、藍色原子筆或鋼筆填寫。',
        10, indent=0.3, sb=0, sa=4)

    # 第一部分：申論題 Q1 & Q2（各 20 分）
    blk(doc, '■ 第一部分：教育理念申論題（共 2 題，每題 20 分，共 40 分）',
        11, bold=True, sb=8, sa=4)
    blk(doc, '※ 作答內容請勿書寫服務單位、姓名或可供辨識身分之個人資料。', 10, indent=0.3, sb=0, sa=6)

    add_essay_space(doc, '第一題（20 分）', n_lines=24)

    doc.add_page_break()
    add_essay_space(doc, '第二題（20 分）', n_lines=28)

    # 第二部分：國語文選擇題 Q1–10
    doc.add_page_break()
    add_mc_answer_grid(doc,
        '■ 第二部分：國語文基本能力選擇題（Q1–10，每題 3 分，共 30 分）', 1, 10)

    blk(doc, '', sb=10, sa=0)

    # 第三部分：數學填充題 Q1–3
    add_math_answer_boxes(
        doc,
        '■ 第三部分：數學基本能力填充題（共 3 題，每題 10 分，共 30 分）',
        ['第 1 題', '第 2 題', '第 3 題']
    )

    blk(doc, '', sb=20, sa=0)
    blk(doc, '— 答案卷結束 —', 11, bold=True, center=True)

    doc.save(path)
    print(f'✓ 第二節答案卷：{path}')


if __name__ == '__main__':
    build_answer1(r'D:\D\onedrive\文件\新版答案卷_第一節_教師意向暨雙語教學知能.docx')
    build_answer2(r'D:\D\onedrive\文件\新版答案卷_第二節_教師意向測驗暨國語文數學.docx')
    print('\n✅ 兩份答案卷均已產生完畢！')
