import docx, sys
sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document(r'D:\D\114設備組\第 66 屆雲林縣科展.docx')
lines = []
for p in doc.paragraphs:
    if p.text.strip():
        lines.append(p.text)
for table in doc.tables:
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if any(cells):
            lines.append('\t'.join(cells))
out = '\n'.join(lines)
with open(r'C:\Users\user\allen\_docx_out.txt', 'w', encoding='utf-8') as f:
    f.write(out)
print('done')
