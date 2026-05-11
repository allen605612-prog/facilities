"""
產生兩份標準答案暨評分要點（docx）
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "標楷體"


# ── 基礎工具 ────────────────────────────────────────────────

def set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    for sec in doc.sections:
        sec.top_margin    = Cm(top)
        sec.bottom_margin = Cm(bottom)
        sec.left_margin   = Cm(left)
        sec.right_margin  = Cm(right)


def _run(para, text, size=12, bold=False, color=None):
    r = para.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    rPr = r._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT)
    rPr.insert(0, rFonts)
    return r


def blk(doc, text='', size=12, bold=False, center=False,
        indent=0.0, sb=2, sa=2, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        _run(p, text, size=size, bold=bold, color=color)
    return p


def section_title(doc, text):
    """深色底框的大節標題"""
    p = blk(doc, text, size=12, bold=True, sb=10, sa=4)
    # 段落底色（使用段落格式底線代替，不依賴 XML 底色以保持相容性）
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def bullet(doc, text, indent=0.8, size=11):
    blk(doc, '• ' + text, size=size, indent=indent, sb=1, sa=1)


# ── MC 答案表 ──────────────────────────────────────────────

def add_mc_answer_table(doc, answers_dict):
    """answers_dict: {題號: '答案字母', ...}"""
    nums = sorted(answers_dict.keys())
    n = len(nums)
    # 每列最多 10 題
    chunk = 10
    for batch_start in range(0, n, chunk):
        batch = nums[batch_start: batch_start + chunk]
        tbl = doc.add_table(rows=3, cols=len(batch) + 1)
        tbl.style = 'Table Grid'

        # 標題欄
        def cell_set(row_i, col_i, text, bold=False, size=11):
            cell = tbl.cell(row_i, col_i)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            _run(p, text, size=size, bold=bold)

        cell_set(0, 0, '題　號', bold=True)
        cell_set(1, 0, '正確答案', bold=True)
        cell_set(2, 0, '配　分', bold=True)

        for i, num in enumerate(batch):
            cell_set(0, i + 1, str(num))
            cell_set(1, i + 1, answers_dict[num], bold=True)
            cell_set(2, i + 1, '')   # 留空，教師自填

        for row in tbl.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trH = OxmlElement('w:trHeight')
            trH.set(qn('w:val'), str(int(Cm(0.85).emu / 914.4)))
            trH.set(qn('w:hRule'), 'exact')
            trPr.append(trH)

        doc.add_paragraph()


# ─────────────────────────────────────────────────────────────
# ★ 第一節標準答案
# ─────────────────────────────────────────────────────────────
def build_key1(path):
    doc = Document()
    set_margins(doc)

    blk(doc, '國立臺中教育大學', 15, bold=True, center=True, sb=0, sa=2)
    blk(doc, '國民小學教師及加註雙語教學次專長學士後教育學分班（B 班）招生考試', 13, bold=True, center=True, sb=0, sa=2)
    blk(doc, '【第一節】教師意向（教育理念）暨雙語教學知能測驗　標準答案暨評分要點', 12, bold=True, center=True, sb=0, sa=6)
    blk(doc, '（本文件僅供閱卷教師參考，請勿外流）', 10, center=True, sb=0, sa=8)

    # ── 第一部分：教育理念選擇題 ─────────────────────────
    section_title(doc, '■ 第一部分：教育理念選擇題（Q1–10，每題 4 分，共 40 分）')

    mc1 = {
        1: 'C',  # 學習共同體起源→日本
        2: 'B',  # Peters 合價值性→值得學習且有正面意義
        3: 'C',  # 核心素養面向→自主行動
        4: 'B',  # ZPD→提供鷹架支持
        5: 'A',  # 觀察學習四歷程→注意→保留→再生→動機
        6: 'B',  # 差異化教學→不同難度任務
        7: 'D',  # Bloom→評鑑（Evaluate）
        8: 'B',  # Gardner→音樂智能
        9: 'C',  # PERMA R→人際關係
        10: 'B', # OCB→主動協助不計報酬
    }
    add_mc_answer_table(doc, mc1)

    # 解析
    notes1 = [
        '1. (C) 學習共同體由日本學者佐藤學提出，強調合作學習與開放課堂。',
        '2. (B) Peters 合價值性：教育內容須值得學習，對學習者具有正面意義。',
        '3. (C) 十二年國教核心素養三面向：自主行動、溝通互動、社會參與。',
        '4. (B) ZPD 教學意涵：由教師提供鷹架（scaffolding）協助學習者超越現有能力。',
        '5. (A) Bandura 觀察學習四歷程：注意（attention）→保留（retention）→再生（reproduction）→動機（motivation）。',
        '6. (B) 差異化教學核心：依學生準備度設計不同難度任務。',
        '7. (D) Bloom 修訂版最高層次為「創造」，但「評斷優缺點」屬「評鑑（Evaluate）」。',
        '8. (B) Gardner 多元智能：辨別音調/旋律/節奏→音樂智能（Musical Intelligence）。',
        '9. (C) PERMA：P-正向情緒、E-投入、R-人際關係、M-意義、A-成就。',
        '10. (B) 組織公民行為（OCB）：超越職責、自發利他、不計個人回報。',
    ]
    blk(doc, '【選擇題解析】', 11, bold=True, sb=8, sa=4)
    for n in notes1:
        bullet(doc, n)

    # ── 第二部分：英文知能選擇題 ────────────────────────
    doc.add_paragraph()
    section_title(doc, '■ 第二部分：雙語教學英文知能選擇題（Q11–20，每題 4 分，共 40 分）')

    mc2 = {
        11: 'B',  # additive bilingualism
        12: 'C',  # CLIL 4Cs Cognition
        13: 'C',  # Krashen i+1
        14: 'C',  # CALP example
        15: 'B',  # translanguaging
        16: 'C',  # outer circle: India, Singapore, Nigeria
        17: 'C',  # dual language immersion
        18: 'C',  # language through learning
        19: 'B',  # bilingual → executive control
        20: 'C',  # debate → varies by factors
    }
    add_mc_answer_table(doc, mc2)

    notes2 = [
        '11. (B) Additive bilingualism：第二語言豐富原有語言庫而不取代母語；對比 subtractive bilingualism（第二語言取代母語）。',
        '12. (C) CLIL 4Cs：Content（內容）、Communication（溝通）、Cognition（認知/思維）、Culture（文化）；Cognition 指發展高階思考能力。',
        '13. (C) Krashen Input Hypothesis：可理解輸入須稍高於學習者現有程度（i+1），才能促進習得。',
        '14. (C) CALP（Cognitive Academic Language Proficiency）：脫離情境脈絡、認知需求高的學術語言，如撰寫分析文章。',
        '15. (B) Translanguaging：視雙語者完整語言資源為學習工具，而非問題。',
        '16. (C) Kachru 三圓圈模型：outer circle（外圈）為英語在原非英語母語國被制度化使用者，如印度、新加坡、奈及利亞。',
        '17. (C) Dual language immersion（雙語沉浸）：目標為雙語、雙識字，同時服務多數與少數語言學生。',
        '18. (C) Language triptych：(a) language of learning=學科專業詞彙；(b) language for learning=課堂管理語言；(c) language through learning=學習過程中自然浮現的語言。',
        '19. (B) 研究指出雙語者需不斷切換、抑制語言，因而強化執行功能（executive functions），包括注意力與任務切換能力。',
        '20. (C) 文章最後指出認知優勢可能因雙語程度、使用頻率與社經背景等因素而有所差異，並非一律成立。',
    ]
    blk(doc, '【選擇題解析】', 11, bold=True, sb=8, sa=4)
    for n in notes2:
        bullet(doc, n)

    # ── 第三部分：申論題 ────────────────────────────────
    doc.add_page_break()
    section_title(doc, '■ 第三部分：申論題（第一題 20 分）')

    blk(doc, '題目：試就正反兩方觀點，評析雙語教育政策對國民小學學生學習與文化認同之影響，並提出推展關鍵原則。', 11, indent=0.5, sb=4, sa=6)

    blk(doc, '【評分說明】', 11, bold=True, sb=4, sa=2)
    blk(doc, '評分項目（各分項加總）', 11, indent=0.5, sb=2, sa=2)

    rubric1 = [
        ('正方觀點（6 分）',
         ['提升英語能力，增強學生國際競爭力（2 分）',
          '培養雙語人才，符合 2030 雙語國家政策目標（2 分）',
          '豐富學習情境，激發學生學習動機（2 分）']),
        ('反方觀點（6 分）',
         ['可能壓縮母語與本土文化學習空間，影響文化認同（2 分）',
          '城鄉師資與資源分配不均，恐擴大教育落差（2 分）',
          '師資培訓不足，若強制推行恐影響教學品質（2 分）']),
        ('推展關鍵原則（6 分）',
         ['循序漸進，依學校條件彈性調整，不一刀切（2 分）',
          '強化師資培訓與行政資源支援（2 分）',
          '母語與英語並重，融入本土文化元素（2 分）']),
        ('論述邏輯與文字表達（2 分）',
         ['論述有條理、舉例具體、立場明確者得滿分']),
    ]

    for title, items in rubric1:
        blk(doc, title, 11, bold=True, indent=0.5, sb=6, sa=2)
        for item in items:
            bullet(doc, item, indent=1.2)

    blk(doc, '【評分說明補充】', 11, bold=True, sb=8, sa=2)
    blk(doc, '正反觀點均須論及方可得分；僅陳述單方觀點者正反部分最高各得 3 分。'
        '推展原則若僅列點未說明，最高得 3 分。文字超過 400 字且論述完整者加分不超過 2 分（含於總分內）。',
        10.5, indent=0.5, sb=0, sa=4)

    # 頁尾
    blk(doc, '', sb=20)
    blk(doc, '— 第一節標準答案暨評分要點 結束 —', 11, bold=True, center=True)

    doc.save(path)
    print(f'✓ 第一節標準答案：{path}')


# ─────────────────────────────────────────────────────────────
# ★ 第二節標準答案
# ─────────────────────────────────────────────────────────────
def build_key2(path):
    doc = Document()
    set_margins(doc)

    blk(doc, '國立臺中教育大學', 15, bold=True, center=True, sb=0, sa=2)
    blk(doc, '國民小學教師及加註雙語教學次專長學士後教育學分班（B 班）招生考試', 13, bold=True, center=True, sb=0, sa=2)
    blk(doc, '【第二節】教師意向測驗（國語文暨數學基本能力）　標準答案暨評分要點', 12, bold=True, center=True, sb=0, sa=6)
    blk(doc, '（本文件僅供閱卷教師參考，請勿外流）', 10, center=True, sb=0, sa=8)

    # ── 第一部分：申論題 ─────────────────────────────────
    section_title(doc, '■ 第一部分：教育理念申論題（共 2 題，每題 20 分，共 40 分）')

    # 申論 Q1
    blk(doc, '第一題（20 分）', 11, bold=True, sb=6, sa=4)
    blk(doc, '題目：試從「親師合作」與「社區資源連結」說明教師如何善用外部夥伴關係，促進學生全人發展。', 11, indent=0.5, sb=2, sa=6)

    rubric_e1 = [
        ('親師合作面向（8 分）',
         ['定期召開親師懇談會，建立雙向溝通機制（2 分）',
          '建立家長志工制度，邀請家長協同教學或班級服務（2 分）',
          '善用數位平台（如 ClassDojo、School+）維持日常聯繫（2 分）',
          '邀請家長參與學校課程設計或決策，發揮共同治理精神（2 分）']),
        ('社區資源連結面向（8 分）',
         ['與社區機構（圖書館、文化中心、企業）建立合作關係（2 分）',
          '邀請社區專業人士入班分享，豐富學生學習視野（2 分）',
          '結合在地文化、產業資源設計課程，落實在地化教育（2 分）',
          '安排學生走入社區進行服務學習或戶外教學（2 分）']),
        ('論述邏輯與文字表達（4 分）',
         ['舉例具體、論述有層次、兩面向均完整說明者得滿分']),
    ]
    for title, items in rubric_e1:
        blk(doc, title, 11, bold=True, indent=0.5, sb=6, sa=2)
        for item in items:
            bullet(doc, item, indent=1.2)

    # 申論 Q2
    doc.add_paragraph()
    blk(doc, '第二題（20 分）', 11, bold=True, sb=6, sa=4)
    blk(doc, '題目：試說明「素養」意涵，並就「真實情境」、「跨領域整合」與「學習遷移」三面向闡述素養導向教學原則。', 11, indent=0.5, sb=2, sa=6)

    rubric_e2 = [
        ('素養意涵（4 分）',
         ['素養是知識、技能與態度/價值觀的整合（2 分）',
          '強調能在真實生活情境中靈活應用、解決問題的能力（2 分）']),
        ('真實情境面向（4 分）',
         ['設計與學生生活緊密連結的學習任務（2 分）',
          '讓學生面對真實問題並提出具體解決方案（2 分）']),
        ('跨領域整合面向（4 分）',
         ['打破單一學科界線，統整多領域知識設計主題課程（2 分）',
          '透過跨領域協作讓學生理解知識的整體關聯（2 分）']),
        ('學習遷移面向（4 分）',
         ['重視深度理解而非表面記憶，使學生能舉一反三（2 分）',
          '設計評量任務要求學生將所學應用於陌生的新情境（2 分）']),
        ('論述邏輯與文字表達（4 分）',
         ['三面向均完整說明、舉例具體、結構清晰者得滿分']),
    ]
    for title, items in rubric_e2:
        blk(doc, title, 11, bold=True, indent=0.5, sb=6, sa=2)
        for item in items:
            bullet(doc, item, indent=1.2)

    # ── 第二部分：國語文選擇題 ─────────────────────────
    doc.add_page_break()
    section_title(doc, '■ 第二部分：國語文基本能力選擇題（Q1–10，每題 3 分，共 30 分）')

    mc3 = {
        1: 'A',   # 出類拔萃、草菅人命 全正確
        2: 'B',   # 頂真、層遞
        3: 'B',   # 力排眾議、堅持己見 全正確
        4: 'C',   # 推敲
        5: 'B',   # 存現句：山腳下散落著幾間農舍
        6: 'D',   # 璧→形聲
        7: 'B',   # 雙關、對偶
        8: 'D',   # 草木 不是聯綿詞
        9: 'B',   # 敬請老師惠賜指正（語用正確）
        10: 'B',  # 靜夜思→思鄉懷遠
    }
    add_mc_answer_table(doc, mc3)

    notes3 = [
        '1. (A) 「萃」（出類拔萃）、「菅」（草菅人命）均正確。'
        '(B) 義無「反」顧（非「返」）；(C) 一「見」鍾情（非「劍」）；(D) 沉「湎」往事（非「緬」）。',

        '2. (B) 頂真：「地利」出現於前句末、後句首；層遞：天時＜地利＜人和，形成遞進關係。',

        '3. (B) 「力排眾議」與「堅持己見」均無錯字。'
        '(A) 不「屈」不撓（非「棄」）；(C) 旁「徵」博引（此處「引徵」有誤）；(D) 立「意」（非「義」）。',

        '4. (C) 唐代詩人賈島在韓愈面前推敲「推」或「敲」，後世以「推敲」比喻斟酌文字。',

        '5. (B) 存現句結構：處所詞語＋動詞＋存在事物（山腳下散落著幾間農舍）。',

        '6. (D) 「璧」：辟（聲符）＋玉（義符），屬形聲字。',

        '7. (B) 「晴/情」諧音雙關；「東邊日出西邊雨」與「道是無晴卻有晴」形成對偶。',

        '8. (D) 聯綿詞兩字不可分解；「草木」中「草」與「木」各有獨立意義，非聯綿詞。'
        '蝴蝶、彷彿、蜘蛛均為聯綿詞。',

        '9. (B) 「拙作」自謙己作、「惠賜指正」請他人給予指導，用語均正確。'
        '(A) 享年用於已逝者，不可用於在世者；(C) 拜讀是己之謙詞，不可要求他人「拜讀」自己作品；'
        '(D) 「拙作」是自謙詞，不可稱對方作品為「拙作」。',

        '10. (B) 〈靜夜思〉藉明月、白霜等意象，抒發客居異鄉的思鄉之情（思鄉懷遠）。',
    ]
    blk(doc, '【選擇題解析】', 11, bold=True, sb=8, sa=4)
    for n in notes3:
        bullet(doc, n)

    # ── 第三部分：數學填充題 ────────────────────────────
    doc.add_page_break()
    section_title(doc, '■ 第三部分：數學基本能力填充題（共 3 題，每題 10 分，共 30 分）')

    math_data = [
        ('第 1 題', '−2',
         '解題：AB 全長 = 10 − (−6) = 16\n'
         'AC : CB = 1 : 3，故 C = A + (1/4)×16 = −6 + 4 = −2'),
        ('第 2 題', '300',
         '解題：a₇ − a₃ = 4d = 27 − 11 = 16，∴ d = 4\n'
         'a₃ = a₁ + 2d → a₁ = 11 − 8 = 3\n'
         'S₁₂ = 12/2 × [2×3 + (12−1)×4] = 6 × (6+44) = 6 × 50 = 300'),
        ('第 3 題', '8',
         '解題：判別式 Δ = (−8)² − 4×2×k = 64 − 8k = 0\n'
         '∴ k = 8　（驗算：2x² − 8x + 8 = 0 → (x−2)² = 0，重根 x = 2）'),
    ]

    for title, ans, sol in math_data:
        blk(doc, title, 11, bold=True, sb=8, sa=2)
        blk(doc, f'標準答案：{ans}', 11, indent=0.5, sb=2, sa=2)
        blk(doc, f'解題過程：', 11, indent=0.5, sb=2, sa=1)
        for line in sol.split('\n'):
            bullet(doc, line.strip(), indent=1.0)
        doc.add_paragraph()

    blk(doc, '【計分說明】各題僅答案填答正確得 10 分；若作答過程部分正確，教師可酌情給 5 分。', 10.5, indent=0.5, sb=4, sa=4)

    # 頁尾
    blk(doc, '', sb=20)
    blk(doc, '— 第二節標準答案暨評分要點 結束 —', 11, bold=True, center=True)

    doc.save(path)
    print(f'✓ 第二節標準答案：{path}')


if __name__ == '__main__':
    build_key1(r'D:\D\onedrive\文件\標準答案_第一節_教師意向暨雙語教學知能.docx')
    build_key2(r'D:\D\onedrive\文件\標準答案_第二節_教師意向測驗暨國語文數學.docx')
    print('\n✅ 兩份標準答案均已產生！')
